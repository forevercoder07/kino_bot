"""
Kinolar ro'yxatini Word (.docx) formatga export qilish moduli.
Ishlatish: from export_films import create_films_docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import tempfile, os, zipfile, io


def _fix_zoom(docx_bytes):
    """python-docx settings.xml dagi zoom xatoligini tuzatish"""
    from lxml import etree
    src = io.BytesIO(docx_bytes)
    dst = io.BytesIO()
    with zipfile.ZipFile(src, 'r') as zin:
        with zipfile.ZipFile(dst, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/settings.xml':
                    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    tree = etree.fromstring(data)
                    for zoom in tree.findall(f'{{{ns}}}zoom'):
                        if f'{{{ns}}}percent' not in zoom.attrib:
                            zoom.set(f'{{{ns}}}percent', '100')
                    data = etree.tostring(tree, xml_declaration=True,
                                         encoding='UTF-8', standalone=True)
                zout.writestr(item, data)
    return dst.getvalue()


def _set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(b)
    # tcBorders: tcW, gridSpan dan keyin
    insert_after = ['tcW', 'gridSpan', 'hMerge', 'vMerge']
    idx = 0
    for i, child in enumerate(tcPr):
        if child.tag.split('}')[1] in insert_after:
            idx = i + 1
    tcPr.insert(idx, tcBorders)


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    # shd: tcBorders dan keyin
    insert_after = ['tcW', 'gridSpan', 'hMerge', 'vMerge', 'tcBorders']
    idx = 0
    for i, child in enumerate(tcPr):
        if child.tag.split('}')[1] in insert_after:
            idx = i + 1
    tcPr.insert(idx, shd)


def _style_cell(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
                color=RGBColor(0x1A, 0x1A, 0x1A), bg="FFFFFF",
                font_size=10, with_borders=True):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if with_borders:
        _set_cell_borders(cell)
    _set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.name = "Arial"


def create_films_docx(films_data: list) -> bytes:
    """
    films_data: [{'name': str, 'code': str, 'parts_count': int}, ...]
    Returns: bytes (docx fayl kontenti)
    """
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    today = datetime.now().strftime("%d.%m.%Y")
    total_parts = sum(f['parts_count'] for f in films_data)

    def add_centered(text, size, color, bold=False, space_after=2):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = "Arial"

    add_centered("KINO VIBE BOT", 18, RGBColor(0x1A, 0x3C, 0x5E), bold=True)
    add_centered("Kinolar ro'yxati", 12, RGBColor(0x55, 0x55, 0x55))
    add_centered(
        f"Sana: {today}     |     Jami: {len(films_data)} ta kino",
        10, RGBColor(0x88, 0x88, 0x88), space_after=8
    )

    # Jadval
    table = doc.add_table(rows=1, cols=4)
    col_widths = [Cm(1.2), Cm(8.5), Cm(2.8), Cm(2.8)]
    for i, w in enumerate(col_widths):
        for c in table.columns[i].cells:
            c.width = w

    # Header qatori
    hdr = table.rows[0]
    hdr.height = Cm(0.9)
    headers = [
        ("№",          WD_ALIGN_PARAGRAPH.CENTER),
        ("Kino nomi",  WD_ALIGN_PARAGRAPH.LEFT),
        ("Kod",        WD_ALIGN_PARAGRAPH.CENTER),
        ("Qismlar",    WD_ALIGN_PARAGRAPH.CENTER),
    ]
    for i, (txt, aln) in enumerate(headers):
        _style_cell(hdr.cells[i], txt, align=aln, bold=True,
                    color=RGBColor(0xFF, 0xFF, 0xFF), bg="1A3C5E", font_size=11)

    # Ma'lumot qatorlari
    for idx, film in enumerate(films_data):
        row = table.add_row()
        row.height = Cm(0.75)
        bg = "FFFFFF" if idx % 2 == 0 else "EBF3FB"
        _style_cell(row.cells[0], str(idx + 1), WD_ALIGN_PARAGRAPH.CENTER,
                    color=RGBColor(0x55, 0x55, 0x55), bg=bg)
        _style_cell(row.cells[1], film['name'], WD_ALIGN_PARAGRAPH.LEFT,
                    bold=True, color=RGBColor(0x1A, 0x1A, 0x1A), bg=bg)
        _style_cell(row.cells[2], str(film['code']), WD_ALIGN_PARAGRAPH.CENTER,
                    color=RGBColor(0x1A, 0x3C, 0x5E), bg=bg)
        _style_cell(row.cells[3], str(film['parts_count']), WD_ALIGN_PARAGRAPH.CENTER,
                    color=RGBColor(0x1A, 0x1A, 0x1A), bg=bg)

    # Footer qatori (merged)
    frow = table.add_row()
    frow.height = Cm(0.8)
    m1 = frow.cells[0].merge(frow.cells[1])
    _style_cell(m1, f"Jami kinolar: {len(films_data)} ta",
                WD_ALIGN_PARAGRAPH.RIGHT, bold=True,
                color=RGBColor(0x1A, 0x3C, 0x5E), bg="F0F4F8", with_borders=False)
    m2 = frow.cells[2].merge(frow.cells[3])
    _style_cell(m2, f"Jami qismlar: {total_parts} ta",
                WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                color=RGBColor(0x1A, 0x3C, 0x5E), bg="F0F4F8", with_borders=False)

    # Pastki eslatma
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(6)
    rn = note.add_run("@kino_vibe_bot tomonidan yaratildi")
    rn.italic = True
    rn.font.size = Pt(9)
    rn.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    rn.font.name = "Arial"

    # Bytes sifatida saqlash
    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)
    with open(tmp, 'rb') as f:
        data = f.read()
    os.remove(tmp)

    return _fix_zoom(data)
